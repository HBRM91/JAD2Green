"""
Bilan Carbone DOCX report renderer.

§0.11 — Report is built exclusively from report_snapshot aggregates.
         Raw activity_facts, provenance, and document content are NEVER
         passed to this module.
§0.12 — AI transparency disclosure block is always included.
"""

from __future__ import annotations

import io
from datetime import datetime
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .narrative import DISCLOSURE_BLOCK


def _dec(val) -> str:
    """Format Decimal or numeric value as string with 2 dp."""
    try:
        return f"{float(Decimal(str(val))):.2f}"
    except Exception:
        return "0.00"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_png(doc: Document, png_bytes: bytes, width_inches: float = 5.5) -> None:
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))


def render_bilan_carbone_docx(
    snapshot: dict,
    project_name: str,
    client_name: str,
    methodology_name: str,
    narrative: dict,
    charts: dict,
) -> bytes:
    """
    Build a Bilan Carbone DOCX report.

    Args:
        snapshot: report_snapshot row (aggregate fields only — §0.11)
        project_name: from projects table
        client_name: from clients table
        methodology_name: from methodologies table
        narrative: {exec_summary, key_findings, recommendations} from LLM
        charts: {scope_pie: bytes, category_bar: bytes, scope2_bar: bytes}

    Returns:
        DOCX file as bytes, ready for download or Google export.
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── 1. Title page ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BILAN CARBONE®")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x38)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Rapport d'émissions de gaz à effet de serre\n").font.size = Pt(14)
    sub.add_run(f"Projet : {project_name}\n").font.size = Pt(13)
    sub.add_run(f"Client : {client_name}\n").font.size = Pt(12)
    sub.add_run(f"Année de référence : {snapshot.get('reporting_year', '—')}\n").font.size = Pt(12)
    sub.add_run(f"Méthode : {methodology_name}\n").font.size = Pt(12)
    sub.add_run(f"Base GWP : {snapshot.get('gwp_basis', 'AR5')}\n").font.size = Pt(11)
    sub.add_run(
        f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    ).font.size = Pt(10)

    doc.add_page_break()

    # ── 2. AI transparency disclosure (§0.12 — always present) ───────────
    _heading(doc, "Avis de transparence IA", level=1)
    p = doc.add_paragraph(DISCLOSURE_BLOCK)
    p.style.font.size = Pt(9)

    doc.add_page_break()

    # ── 3. Executive summary ──────────────────────────────────────────────
    _heading(doc, "Synthèse exécutive", level=1)
    doc.add_paragraph(narrative.get("exec_summary", ""))

    doc.add_paragraph()
    _heading(doc, "Constats clés", level=2)
    for finding in narrative.get("key_findings", []):
        doc.add_paragraph(finding, style="List Bullet")

    doc.add_page_break()

    # ── 4. Résultats par scope ────────────────────────────────────────────
    _heading(doc, "Résultats des émissions", level=1)
    totals = snapshot.get("totals_co2e", {})
    uncertainty = snapshot.get("uncertainty", {})

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Scope", "tCO₂e", "Incertitude ±", "Note"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    def _unc_pct(key: str) -> str:
        u = uncertainty.get(key, {})
        frac = u.get("fraction", "0")
        try:
            return f"{float(Decimal(str(frac))) * 100:.0f}%"
        except Exception:
            return "—"

    rows = [
        ("Scope 1 — Émissions directes", totals.get("scope1", "0"), _unc_pct("scope1"), ""),
        (
            "Scope 2 — Location-based",
            snapshot.get("scope2_location_t") or totals.get("scope2_location", "0"),
            _unc_pct("scope2_location"),
            "Méthode réseau",
        ),
        (
            "Scope 2 — Market-based",
            snapshot.get("scope2_market_t") or totals.get("scope2_market", "0"),
            _unc_pct("scope2_market"),
            "Méthode marché",
        ),
        ("Scope 3 — Indirect", totals.get("scope3", "0"), _unc_pct("scope3"), ""),
        ("TOTAL (S1 + S2 loc + S3)", totals.get("total", "0"), _unc_pct("total"), "Périmètre GHG Protocol"),
    ]
    for label, value, unc, note in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _dec(value)
        row[2].text = unc
        row[3].text = note

    doc.add_paragraph()

    # ── 5. Charts ─────────────────────────────────────────────────────────
    if charts.get("scope_pie"):
        _heading(doc, "Répartition des émissions", level=2)
        _add_png(doc, charts["scope_pie"], width_inches=4.5)
        doc.add_paragraph()

    if charts.get("scope2_bar"):
        _heading(doc, "Scope 2 : Location vs Marché", level=2)
        _add_png(doc, charts["scope2_bar"], width_inches=4.5)
        doc.add_paragraph()

    if charts.get("category_bar"):
        doc.add_page_break()
        _heading(doc, "Émissions par catégorie", level=2)
        _add_png(doc, charts["category_bar"], width_inches=5.5)
        doc.add_paragraph()

    # ── 6. GRI 305 disclosure (when data present) ─────────────────────────
    gri_305 = snapshot.get("gri_305_data")
    if gri_305:
        doc.add_page_break()
        _heading(doc, "GRI 305 — Divulgation des émissions", level=1)
        doc.add_paragraph(
            "Conformément à la norme GRI 305 (Émissions), les divulgations suivantes sont issues "
            "du snapshot de calcul Adrar AI. Ces données sont calculées à partir des faits d'activité "
            "validés par un consultant habilité.\n"
        )
        gri_table = doc.add_table(rows=1, cols=3)
        gri_table.style = "Light Grid Accent 1"
        gri_hdr = gri_table.rows[0].cells
        for i, h in enumerate(["Indicateur GRI", "Valeur (tCO₂e)", "Description"]):
            gri_hdr[i].text = h
            gri_hdr[i].paragraphs[0].runs[0].font.bold = True

        gri_rows = [
            ("GRI 305-1", gri_305.get("305-1", 0), "Émissions directes de GES (Scope 1)"),
            ("GRI 305-2 (loc.)", gri_305.get("305-2-loc", 0), "Émissions indirectes — réseau (location-based)"),
            ("GRI 305-2 (mkt.)", gri_305.get("305-2-mkt", 0), "Émissions indirectes — marché (market-based)"),
            ("GRI 305-3", gri_305.get("305-3", 0), "Autres émissions indirectes (Scope 3)"),
            ("GRI 305 Total (loc.)", gri_305.get("305-total-loc", 0), "Scope 1 + Scope 2 loc. + Scope 3"),
            ("GRI 305 Total (mkt.)", gri_305.get("305-total-mkt", 0), "Scope 1 + Scope 2 mkt. + Scope 3"),
        ]
        for indicator, value, desc in gri_rows:
            row = gri_table.add_row().cells
            row[0].text = indicator
            row[1].text = _dec(value)
            row[2].text = desc

        doc.add_paragraph()
        doc.add_paragraph(
            "Note : La divulgation GRI 305-4 (Intensité des émissions) et GRI 305-5 (Réductions) "
            "nécessitent un indicateur de dénominateur et des données historiques respectivement. "
            "Ces divulgations doivent être complétées par le consultant."
        ).style.font.size = Pt(9)

    # ── 6b. NDC Morocco alignment (when data present) ────────────────────
    ndc = snapshot.get("ndc_alignment")
    if ndc and ndc.get("baseline_year"):
        doc.add_page_break()
        _heading(doc, "Alignement NDC Maroc 2030", level=1)
        doc.add_paragraph(
            "L'objectif climatique du Maroc (NDC 2030) prévoit une réduction de 45,5% des émissions "
            "par rapport au scénario de référence (BAU). Le tableau ci-dessous présente l'état d'avancement "
            f"du projet par rapport à cette cible, en utilisant {ndc.get('baseline_year')} comme année de référence."
        )
        doc.add_paragraph()

        ndc_table = doc.add_table(rows=1, cols=2)
        ndc_table.style = "Light Grid Accent 1"
        ndc_hdr = ndc_table.rows[0].cells
        ndc_hdr[0].text = "Indicateur"
        ndc_hdr[0].paragraphs[0].runs[0].font.bold = True
        ndc_hdr[1].text = "Valeur"
        ndc_hdr[1].paragraphs[0].runs[0].font.bold = True

        ndc_rows_data = [
            ("Année de référence", str(ndc.get("baseline_year", "—"))),
            ("Émissions de référence (tCO₂e)", _dec(ndc.get("baseline_emissions", 0))),
            ("Émissions actuelles (tCO₂e)", _dec(ndc.get("current_emissions", 0))),
            ("Cible NDC 2030 (tCO₂e)", _dec(ndc.get("target_emissions", 0))),
            ("Réduction accomplie (tCO₂e)", _dec(ndc.get("reduction_achieved", 0))),
            ("Progression vers l'objectif NDC", f"{ndc.get('progress_pct', 0):.1f}%"),
            ("En bonne voie (on track)", "Oui" if ndc.get("on_track") else "Non" if ndc.get("on_track") is False else "—"),
        ]
        for label, value in ndc_rows_data:
            row = ndc_table.add_row().cells
            row[0].text = label
            row[1].text = value

        doc.add_paragraph()
        doc.add_paragraph(
            "Source : Contribution Déterminée au niveau National (NDC) du Maroc, 2021. "
            "Cible conditionnelle : −45,5% par rapport au BAU d'ici 2030."
        ).style.font.size = Pt(9)

    # ── 7. Recommandations ────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Recommandations", level=1)
    for rec in narrative.get("recommendations", []):
        doc.add_paragraph(rec, style="List Bullet")

    # ── 8. Annexe : empreinte du calcul ──────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Annexe — Empreinte du calcul", level=1)
    doc.add_paragraph(
        f"Hash d'état (state_hash) : {snapshot.get('state_hash', '—')}\n"
        f"Versions des jeux de facteurs : {', '.join(snapshot.get('factor_set_versions', []))}\n"
        f"Horizon GWP : 100 ans — Base : {snapshot.get('gwp_basis', 'AR5')}\n"
        f"L'incertitude est calculée séparément des totaux conformément aux invariants §0.2."
    )

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
